import time
import docx
import os
import keyboard
from selenium.webdriver.common.keys import Keys
import prefect
from docx import Document
from newspaper import Article
import newspaper
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from seleniumbase import BaseCase
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from prefect import Flow, task, flow
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
from keyboard import press

# Abrimos las variables para las secciones de tecnologia de todos los noticieros

url_parafrasis = 'https://www.parafraseartextos.net/'
url_Milenio = 'https://www.milenio.com/tecnologia'
url_ElUniversal = 'https://www.eluniversal.com.mx/tag/tecnologias-de-la-informacion-y-comunicacion'
url_ElEconomista = 'https://www.eleconomista.com.mx/seccion/tecnologia/'
url_Gaceta = 'http://www.gaceta.udg.mx/'

# array para guardar todos los links
urls = []
# array para guardar titulo de la nota, texto de la nota y el link de la fuente
nota = []
# Aquellas notas que pasen las pruebas seran almacenadas en esta lista
urls_usados = []

#Variable que guarda el articulo



# Funcion para extraer el titulo y el texto de cada pagina

@task
def extraerTextoArticulo(url):
    # Configuramostodo para que el articulo este en español
    mi_Articulo = Article(url, language="es")
    # Ponemos en el sistema el articulo
    mi_Articulo.download()
    # lo leemos
    mi_Articulo.parse()
    #La pagina en la que parafraseamos nos pide un minimo y un maximo de caraacteres
    #para cumplir con las caracteristicas del texto hacemos esta comparacion
    #si no cumple los requisitos eliminamos la nota y procedemos a la siguiente

    if(len(mi_Articulo.text)<300 or len(mi_Articulo.text)>4999):
        return
    #Cuando pasa el filtro pasamos a la lista de nota el titulo en mayuscula, el texto y el url
    nota.append(mi_Articulo.title.upper())
    nota.append(url)
    #extarer articulo

    #Llamamos a la funcion para que abra la pagina y edite el texto
    #parafrasearLink(mi_Articulo.text)


    #### agregamostodo a un txt para que no se pierda
    f = open("nSinParafrasis.txt", "x")  # crear
    f.close()

    f = open("nSinParafrasis.txt", "a+")  # Leer + Append (agregar al final)
    f.write(mi_Articulo.title.upper())
    f.write(mi_Articulo.text)
    f.write(url)
    f.close()

    return(mi_Articulo.text)

    pass
#https://parafrasist.com  pagina para arafrasis sin revision de bot
# https://www.parafraseartextos.net/ segunda opcion y funcional
#funcion para traer los links de tecnologia de la pagina


@task
def traerLosLinks():
    #bandera para reventar el for de las busquedas
    tronar = 0
    #El memoize articles es para borrar los articulos guardados entonces simepre nos va a dar los mas nuevos
    milenio = newspaper.build(url_Milenio, memoize_articles=False)
    universal = newspaper.build(url_ElUniversal, memoize_articles=False)
    economista = newspaper.build(url_ElEconomista, memoize_articles=False)
    gaceta = newspaper.build(url_Gaceta, memoize_articles=False)
    #Fors para obtener los links de las respectivas paginas

    for articulo in gaceta.articles:
        tronar += 1
        print(articulo.url)
        urls.append(articulo.url)
        if(tronar == 3):
            tronar = 0
            break


    for articulo in economista.articles:
        tronar += 1
        print(articulo.url)
        urls.append(articulo.url)
        if(tronar == 3):
            tronar = 0
            break

    for articulo in milenio.articles:
        tronar += 1
        print(articulo.url)
        urls.append(articulo.url)
        if(tronar == 3):
            tronar = 0
            break

    """for articulo in universal.articles:
        tronar += 1
        print(articulo.url)
        urls.append(articulo.url)
        if(tronar == 2):
            tronar = 0
            break"""""
    pass

# Seccion de escribir y guardar en el documento word

@task
def writeonDoc():
    # Especificamos en que direccion esta el documento al que vamos a escribir
    doc = Document('D:\\Drive\\ilab_TDI\\Notas\\notas_publicar_ILaB.docx')
    # Creamos una variable para poder editar el estilo y las fuentes
    content = doc.paragraphs
    # Leemos la longitud del documento para para escribir al final de este
    longitudDocumento = len(content) - 1

    content_run = content[longitudDocumento].runs
    for i in content_run:
        print(i.text)
    # Comenzamos a darle estilo al texto
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph(nota[0])
    doc.add_paragraph("\n")
    doc.add_paragraph(nota[2])
    doc.add_paragraph("\n")
    doc.add_paragraph(nota[1])
    doc.add_paragraph("\n")
    doc.add_paragraph("\n")
    # Guardamos el save
    doc.save('D:\\Drive\\ilab_TDI\\Notas\\notas_publicar_ILaB.docx')
    # Pasamos a la lsita de links usados el link util
    urls_usados.append(nota[1])
    # Limpiamos la lista para dar paso a la siguiente nota
    nota.clear()
    pass




@task
def parafrasearLink(texto):

    #texto = 'Concepto texto. El texto es la unidad superior de comunicación y de la competencia organizacional del hablante. Su extensión es variable y corresponde a un
    #todocomprensible que tiene una finalidad comunicativa en un contexto dado. El carácter comunicativo, pragmático y estructural permiten su identificación.'

    driver = webdriver.Chrome(executable_path='C:/Users/vicen/Downloads/chromedriver.exe')

    driver.get(url_parafrasis)



    text_area = driver.find_element(By.XPATH, '//*[@id="input_box"]')

    text_area.send_keys(texto)
    time.sleep(2)
    button_click = driver.find_element(By.XPATH,'/html/body/section[1]/div/div/div[2]/div[1]/div[2]/div[3]/button')
    driver.execute_script("arguments[0].click();", button_click)
    time.sleep(60)
    webdriver.ActionChains(driver).send_keys(Keys.ESCAPE).perform()
    #button_click_after = driver.find_element(By.XPATH,'/html/body/div[6]/div/div[3]/button[1]')#Podemos hacer que pique escape la compu en lugar de boton
    #driver.execute_script("arguments[0].click();", button_click_after)
    text = driver.find_element(By.XPATH,'/html/body/section[1]/div/div/div[2]/div[2]/div[1]').text
    print(text)
    nota.append(text)

    pass

@task
def imprimirLinksUsados():
    print("Estos son los URLs usados")
    for links in urls_usados:
        print(links)
    return


# Aqui comienza el "main"
@flow(name="parafraseo")
def proceso():
    traerLosLinks()
    for links in urls:
        texto = extraerTextoArticulo(links)
        parafrasearLink(texto)
        writeonDoc()
        if(len(urls_usados)>=5):
            break
    imprimirLinksUsados()

if __name__ == "__main__":
    proceso()

